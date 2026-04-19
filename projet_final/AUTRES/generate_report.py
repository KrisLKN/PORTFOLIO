import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_report():
    doc = Document()
    
    # --- STYLES ---
    style_title = doc.styles['Title']
    font_title = style_title.font
    font_title.name = 'Arial'
    font_title.size = Pt(24)
    font_title.color.rgb = RGBColor(0, 51, 102)

    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Arial'
    font_h1.size = Pt(16)
    font_h1.color.rgb = RGBColor(0, 102, 204)

    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Arial'
    font_h2.size = Pt(14)
    font_h2.color.rgb = RGBColor(0, 153, 204)

    # --- TITLE PAGE ---
    title = doc.add_paragraph("RAPPORT DE PROJET FINAL", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Mise en place d'un Système d'Information Décisionnel (Data Warehouse)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].bold = True
    
    doc.add_paragraph("\n\n\n")
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("SAÉ 4 - Bases de données avancées et Informatique Décisionnelle\n\n\n")
    run.font.size = Pt(14)
    run.bold = True
    
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run("Présenté par :\n").bold = True
    authors.add_run("Kris LOKOUN & Abel THYS\n")
    authors.add_run("BUT SD S4")
    
    doc.add_page_break()

    # --- CONTENT ---
    
    doc.add_heading("I. Introduction et Contexte du Projet", level=1)
    doc.add_paragraph(
        "Ce rapport détaille la conception, le développement et le déploiement d'un Data Warehouse (DWH) complet "
        "dans le cadre de la SAÉ 4. Le projet simule la gestion des données de vente d'une entreprise "
        "distribuant des pistes musicales. L'objectif était de consolider les données provenant de deux sources indépendantes "
        "pour permettre une analyse décisionnelle robuste."
    )
    doc.add_paragraph(
        "Ce projet a marqué une évolution technologique majeure par rapport à nos travaux précédents, passant "
        "d'un écosystème Oracle à un écosystème entièrement Microsoft (SQL Server, SSIS, SSMS). Nous avons dû intégrer :\n"
        "• Une base source Online (magasin en ligne type Chinook).\n"
        "• Une base source Magasin physique (initialement issue d'Oracle et migrée vers SQL Server)."
    )

    doc.add_heading("II. Architecture et Modélisation Décisionnelle", level=1)
    doc.add_heading("1. Architecture globale", level=2)
    doc.add_paragraph(
        "Le flux de données suit une architecture classique en BI :\n"
        "1. Extraction des données des sources hétérogènes.\n"
        "2. Chargement en Staging Area (ODS) pour unifier les formats.\n"
        "3. Intégration dans le Data Warehouse (DWH) structuré pour l'analyse OLAP."
    )

    doc.add_heading("2. Modèle en Étoile (Star Schema)", level=2)
    doc.add_paragraph(
        "Notre DWH repose sur un schéma en étoile très dénormalisé, optimisant les temps de réponse des futures requêtes analytiques. "
        "Le cœur du modèle est la table FACT_VENTES qui stocke les métriques quantitatives (Quantité, Prix Unitaire, Total)."
    )
    doc.add_paragraph("Autour gravitent 5 tables de dimensions :")
    doc.add_paragraph("• DIM_CUSTOMER : les clients", style='List Bullet')
    doc.add_paragraph("• DIM_EMPLOYEE : les vendeurs", style='List Bullet')
    doc.add_paragraph("• DIM_CHANNEL : les canaux de vente (Online/Store)", style='List Bullet')
    doc.add_paragraph("• DIM_DATE : l'axe temporel (jour, mois, année, trimestre)", style='List Bullet')
    doc.add_paragraph("• DIM_TRACK : les pistes musicales vendues (gérant l'historique de prix)", style='List Bullet')
    doc.add_paragraph(
        "La pierre angulaire de cette modélisation est la rupture avec les clés transactionnelles : chaque dimension "
        "dispose de sa propre Clé Technique auto-incrémentée (TK_...) qui remplace la Clé Naturelle (NK_...). "
        "Cela garantit que l'intégrité du DWH ne dépend pas des systèmes sources."
    )

    doc.add_heading("III. Processus d'Intégration ETL (SSIS)", level=1)
    doc.add_paragraph(
        "Le rapatriement et la transformation des données s'effectuent via SQL Server Integration Services (SSIS). "
        "Nous avons conçu plus d'une trentaine de packages pour encadrer ce flux."
    )

    doc.add_heading("1. Historisation des prix : Le SCD Type 2 (DIM_TRACK)", level=2)
    doc.add_paragraph(
        "Un défi technique majeur de cette SAÉ a été la gestion de l'évolution du prix d'une piste musicale dans le temps. "
        "Pour cela, nous avons implémenté un Slowly Changing Dimension (SCD) de Type 2."
    )
    doc.add_paragraph(
        "Mécanisme mis en place : \n"
        "• Les attributs comme le nom de l'artiste ou l'album sont définis en modifications simples (Attributs de modification).\n"
        "• Le prix unitaire (UNIT_PRICE) est un Attribut d'historique. Si le prix change dans la source, SSIS procède à deux actions :\n"
        "   1) Désactivation de l'ancienne version : le champ ACTIF passe à 0, et la colonne DATE_FIN est alimentée via une Commande OLE DB (avec récupération de GETDATE()).\n"
        "   2) Insertion de la nouvelle version : une nouvelle ligne apparaît pour la même piste (NK_ID_PISTE), avec le nouveau prix, ACTIF = 1, et une nouvelle DATE_DEBUT."
    )
    doc.add_paragraph(
        "Cette implémentation permet par exemple de calculer un panier de vente de l'année précédente avec le prix exact "
        "qui était appliqué lors de l'achat, évitant toute corruption comptable."
    )

    doc.add_heading("2. Alimentation de la Table de Faits (FACT_VENTES)", level=2)
    doc.add_paragraph(
        "La table de faits consolide l'ensemble des données. "
        "Le package SSIS procède d'abord à la fusion (UNION ALL) des flux Online et Store. Ensuite, 4 composants Lookup succéssifs "
        "sont utilisés pour récupérer les Clés Techniques (TK) :"
    )
    doc.add_paragraph("• LKP_CLIENT croise le NK_ID_CLIENT avec DIM_CUSTOMER.", style='List Bullet')
    doc.add_paragraph(
        "• LKP_PISTE est primordial : il croise le NK_ID_PISTE avec DIM_TRACK en filtrant impérativement sur ACTIF = 1 "
        "pour associer la vente à la version actuellement en vigueur du produit.", style='List Bullet'
    )
    doc.add_paragraph("• LKP_CANAL et LKP_EMPLOYE finalisent la récupération des contextes.", style='List Bullet')

    doc.add_heading("3. Orchestration Finale", level=2)
    doc.add_paragraph(
        "Le routage global est assuré par un Package Maître (Master_Load_DWH.dtsx). Ce contrôleur :"
    )
    doc.add_paragraph("• Nettoie d'abord les tables cibles (TRUNCATE ou DELETE ciblés).", style='List Bullet')
    doc.add_paragraph("• Exécute en parallèle les dimensions indépendantes (Customer, Employee, Date) pour gagner en performances.", style='List Bullet')
    doc.add_paragraph("• Exécute le chargement dimensionnel gérant le SCD2 (Track).", style='List Bullet')
    doc.add_paragraph("• N'exécute la table de faits (FACT_VENTES) qu'une fois que l'entiereté du référentiel est intégrée (garantie de l'intégrité référentielle FK-PK).", style='List Bullet')

    doc.add_heading("IV. Conclusion et Bilan", level=1)
    doc.add_paragraph(
        "L'interface vers les décideurs est désormais prête. Ce projet nous a permis de franchir un cap dans l'architecture des systèmes. "
        "La conception manuelle d'une historisation SCD2 à travers l'interface SSIS et les requêtes SQL paramétrées fut l'aspect le plus complexe, "
        "mais le plus gratifiant une fois fonctionnel et testé avec succès."
    )
    doc.add_paragraph(
        "Ce Data Warehouse, robuste, propre et tracable, est prêt à héberger des cubes SQL Server Analysis Services (SSAS) ou à "
        "être visualisé dans Power BI."
    )

    doc.add_heading("Bilan Personnel", level=2)
    p_kris = doc.add_paragraph()
    p_kris.add_run("Kris LOKOUN : ").bold = True
    p_kris.add_run("Focus sur la conception SSIS, intégration du SCD Type 2, débogage des flux de données avec les requêtes de test en base, "
                   "gestion des dates de fin complexes et orchestration multi-packages.")
                   
    p_abel = doc.add_paragraph()
    p_abel.add_run("Abel THYS : ").bold = True
    p_abel.add_run("Focus sur les scripts SQL d'initialisation, la migration syntaxique Oracle/T-SQL, et les vérifications des charges "
                   "ainsi que la gestion SQL des environnements.")

    # Save document
    report_path = r"c:\Users\LOKOUN Kris\Desktop\BUT 2\SAE DATA\S4\projet_final\RAPPORT_SAE4_LOKOUN_THYS_FINAL_OFFICIEL.docx"
    doc.save(report_path)
    print(f"Report saved to {report_path}")

if __name__ == '__main__':
    create_report()
