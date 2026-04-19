from docx import Document
from docx.shared import Pt, Inches

def create_word_doc():
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('Plan Détaillé et Préparation du Rapport Final SAE 4', 0)
    
    # Intro
    doc.add_paragraph('Ce document contient le plan détaillé du nouveau rapport à rédiger, la liste exacte des captures d\'écran à prendre pour illustrer vos propos, ainsi qu\'une série de questions auxquelles vous devez répondre. Vos réponses serviront de base pour générer un texte riche, précis et exceptionnel.')
    
    # ---------------------------------------------------------
    # PARTIE 1 : PLAN DU RAPPORT
    # ---------------------------------------------------------
    doc.add_heading('I. Plan Détaillé du Nouveau Rapport', level=1)
    
    plan = [
        ("1. Introduction et Contexte", [
            "Présentation du projet SAE 4 (continuité de la SAE 3).",
            "Choix technologique : passage d'Oracle/ODI à Microsoft SQL Server/SSIS. Pourquoi ce choix (avantages, facilité d'installation, standard du marché).",
            "Organisation du travail en binôme (répartition Kris/Abel, durée du projet)."
        ]),
        ("2. Architecture et Modélisation", [
            "Présentation de l'existant : la base transactionnelle Chinook (MPD).",
            "L'architecture en 3 couches (DSA, ODS, DWH) : explication du rôle de chaque couche avec SQL Server.",
            "Le modèle décisionnel en étoile final (DWH) : définition des faits et dimensions, intégration des Playlists (aplatissement vs flocon)."
        ]),
        ("3. L'Intégration de Nouvelles Données (Magasin)", [
            "Analyse de la deuxième source : la base Magasin (différences, conversion de syntaxe Oracle vers T-SQL).",
            "Le mécanisme de fusion (Union All) dans SSIS pour alimenter les tables de faits à partir de deux flux distincts."
        ]),
        ("4. Les Flux ETL et l'Historisation (SCD Type 2)", [
            "Le fonctionnement global des flux (Truncate pour DSA, Lookup match/no match pour ODS).",
            "La gestion complexe de la dimension DIM_PISTE (Slowly Changing Dimension Type 2) : suivi des modifications de prix, dates de début/fin, et flag Actif.",
            "Résolution des bugs rencontrés (propagation des UPDATE depuis l'ODS, erreur de valeur NULL sur DATE_DEBUT corrigée vie un DEFAULT)."
        ]),
        ("5. Reporting et Pilotage", [
            "Présentation du Dashboard Power BI à 5 onglets.",
            "Les KPI choisis pour la synthèse et l'analyse croisée des ventes (Online vs Magasin)."
        ]),
        ("6. Bilan et Perspectives", [
            "Les leçons apprises (maîtrise d'un nouvel ETL, gestion de l'historique).",
            "Améliorations possibles si plus de temps (automatisation via SQL Server Agent, alertes, etc.)."
        ])
    ]
    
    for title, sub_items in plan:
        doc.add_heading(title, level=2)
        for item in sub_items:
            doc.add_paragraph(item, style='List Bullet')
            
    doc.add_page_break()

    # ---------------------------------------------------------
    # PARTIE 2 : LISTE DES CAPTURES
    # ---------------------------------------------------------
    doc.add_heading("II. Liste des Captures d'Écran à Fournir", level=1)
    doc.add_paragraph("Prenez ces captures, nommez-les correctement et placez-les dans un dossier 'images_rapport'. C'est ce qui donnera une valeur professionnelle maximale à votre rapport.")
    
    captures = [
        "capture_mpd_chinook.png : Le schéma de départ (MPD) de la base Chinook.",
        "capture_etoile_dwh.png : Votre schéma en étoile final pour le DWH.",
        "capture_ssms_bases.png : L'explorateur SSMS montrant vos 7 bases de données créées.",
        "capture_ssms_scd2.png : Une requête SELECT dans DIM_PISTE montrant clairement 2 lignes pour la même piste (l'ancien prix ACTIF=0, le nouveau ACTIF=1).",
        "capture_ssis_master.png : Le Control Flow d'un package maître (RUN_ALL) tout en vert.",
        "capture_ssis_ods.png : Le Data Flow d'un chargement ODS montrant le composant 'Lookup' et le 'OLE DB Command' pour l'UPDATE.",
        "capture_ssis_fusion.png : Le Data Flow de FAIT_VENTE montrant le composant 'Union All' qui mélange le Magasin et Chinook.",
        "capture_ssis_scd2.png : Le flux détaillé de l'historisation de DIM_PISTE.",
        "capture_pbi_synthese.png : Votre onglet Synthèse de Power BI.",
        "capture_pbi_ventes.png : Votre onglet Ventes de Power BI."
    ]
    
    for cap in captures:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_page_break()
    
    # ---------------------------------------------------------
    # PARTIE 3 : QUESTIONNAIRE A REMPLIR
    # ---------------------------------------------------------
    doc.add_heading('III. Questionnaire à Remplir', level=1)
    doc.add_paragraph('Tapez vos réponses détaillées directement sous chaque question dans ce document Word.')
    
    questions = [
        "1. Qu'est-ce qui a été le plus satisfaisant ou le plus difficile en passant d'Oracle (ODI) à l'écosystème Microsoft (SQL Server / SSIS) ?",
        "2. Comment avez-vous réparti concrètement le travail technique avec Abel ? (Qui a fait quoi, et comment vous êtes-vous synchronisés ?)",
        "3. La base \"Magasin\" que vous avez dû intégrer : quelles étaient ses spécificités par rapport à Chinook ? Y a-t-il eu beaucoup de nettoyage ou de renommage à faire avant la fusion ?",
        "4. Pourquoi avez-vous choisi de ramener les informations des Playlists (nombre et noms) directement dans la table DIM_PISTE au lieu de créer une dimension DIM_PLAYLIST séparée en flocon ?",
        "5. Racontez le problème que vous avez eu avec l'historisation (les prix modifiés dans la source n'arrivaient pas dans le DWH). Comment le composant OLE DB Command dans l'ODS a-t-il sauvé la mise ?",
        "6. Expliquez le fameux bug de la \"DATE_DEBUT qui ne peut pas être NULL\" lors de l'insertion dans DIM_PISTE, et comment vous l'avez corrigé en mettant une valeur DEFAULT GETDATE() dans SQL Server.",
        "7. Pour le Dashboard Power BI : quels sont vos 3 indicateurs clés les plus importants dans l'onglet Synthèse ? Avez-vous réussi à croiser les ventes physiques (Magasin) et les ventes en ligne ?"
    ]
    
    for q in questions:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        doc.add_paragraph("Votre réponse : \n\n\n\n")

    # Save format
    doc.save(r'c:\Users\LOKOUN Kris\Desktop\BUT 2\SAE DATA\S4\projet_final\Plan_et_Questions_Rapport_SAE4.docx')

if __name__ == '__main__':
    create_word_doc()
